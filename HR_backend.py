import os
import re
import hashlib
import logging
import tempfile
from datetime import datetime, timedelta
from langchain_openai import ChatOpenAI
from langchain_core.messages import SystemMessage, HumanMessage
import uuid
import json
import boto3
import pdfplumber
from docx import Document
from typing import TypedDict, List, Optional
from langgraph.errors import NodeInterrupt
from pydantic import BaseModel, field_validator, model_validator
from langgraph.graph import StateGraph,START,END
from dotenv import load_dotenv
load_dotenv()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s — %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
logger = logging.getLogger("hiring_agent")

s3 = boto3.client(
    "s3",
    aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
    aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"),
    region_name=os.getenv("AWS_REGION")
)

sqs = boto3.client(
    "sqs",
    aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
    aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"),
    region_name=os.getenv("AWS_REGION")
)

ses = boto3.client(
    "ses",
    aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
    aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"),
    region_name=os.getenv("AWS_REGION"),
)

S3_BUCKET = os.getenv("S3_BUCKET_NAME")
SQS_QUEUE_URL = os.getenv("SQS_QUEUE_URL")
HR_EMAIL = os.getenv("HR_EMAIL", "hr@company.com")

llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.2, model_kwargs={"response_format": {"type": "json_object"}})


def _parse_llm_json(content: str) -> dict:
    content = re.sub(r"^```(?:json)?\s*|\s*```$", "", content.strip())
    return json.loads(content)


class _Base(BaseModel):
    """Coerces None → '' for all str fields so LLM nulls never crash validation."""
    @model_validator(mode="before")
    @classmethod
    def _coerce_none_strings(cls, data):
        if isinstance(data, dict):
            for name, info in cls.model_fields.items():
                if info.annotation is str and data.get(name) is None:
                    data[name] = ""
        return data


class ExperienceRequired(_Base):
    min_years: Optional[int] = None
    max_years: Optional[int] = None
    raw: str = ""

class Skills(_Base):
    must_have: List[str] = []
    nice_to_have: List[str] = []

class ScreeningQuestion(BaseModel):
    question: str
    evaluates: str

class Signals(BaseModel):
    remote_allowed: Optional[bool] = None
    startup_environment: Optional[bool] = None
    client_facing: Optional[bool] = None

class JobProfile(_Base):
    job_title: str = ""
    company: str = ""
    location: str = ""
    employment_type: str = ""
    seniority_level: str = ""
    domain: str = ""
    experience_required: ExperienceRequired = ExperienceRequired()
    skills: Skills = Skills()
    responsibilities: List[str] = []
    qualifications: List[str] = []
    education: List[str] = []
    keywords: List[str] = []
    screening_questions: List[ScreeningQuestion] = []
    signals: Signals = Signals()

class Experience(_Base):
    company: str = ""
    role: str = ""
    duration: str = ""
    description: str = ""

class Education(_Base):
    institution: str = ""
    degree: str = ""
    year: Optional[str] = None

    @field_validator("year", mode="before")
    @classmethod
    def coerce_year(cls, v):
        return str(v) if v is not None else v

class CandidateProfile(_Base):
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    skills: List[str] = []
    experience: List[Experience] = []
    education: List[Education] = []
    certifications: List[str] = []

    @field_validator("skills", mode="before")
    @classmethod
    def coerce_skills(cls, v):
        if not isinstance(v, list):
            return []
        result = []
        for item in v:
            if isinstance(item, str):
                result.append(item)
            elif isinstance(item, dict):
                # e.g. {"skill": "Python", "level": "expert"}
                result.append(
                    item.get("skill") or item.get("name") or item.get("technology") or str(item)
                )
        return result

    @field_validator("experience", mode="before")
    @classmethod
    def coerce_experience(cls, v):
        if not isinstance(v, list):
            return []
        result = []
        for item in v:
            if isinstance(item, str):
                # plain string → put it in description
                result.append({"company": "", "role": "", "duration": "", "description": item})
            elif isinstance(item, dict):
                result.append(item)
        return result

    @field_validator("education", mode="before")
    @classmethod
    def coerce_education(cls, v):
        if not isinstance(v, list):
            return []
        result = []
        for item in v:
            if isinstance(item, str):
                result.append({"institution": item, "degree": "", "year": None})
            elif isinstance(item, dict):
                result.append(item)
        return result

    @field_validator("certifications", mode="before")
    @classmethod
    def coerce_certifications(cls, v):
        if not isinstance(v, list):
            return []
        result = []
        for item in v:
            if isinstance(item, str):
                result.append(item)
            elif isinstance(item, dict):
                # LLM sometimes returns {"name": "...", "year": ...} or {"title": "..."}
                name = item.get("name") or item.get("title") or item.get("certification") or ""
                year = item.get("year")
                result.append(f"{name} ({year})" if name and year else name or str(item))
        return result

class ScoringWeights(BaseModel):
    skills: float = 0.4
    experience: float = 0.3
    education: float = 0.2
    culture_fit: float = 0.1

    def normalized(self) -> dict:
        total = self.skills + self.experience + self.education + self.culture_fit
        return {
            "skills": self.skills / total,
            "experience": self.experience / total,
            "education": self.education / total,
            "culture_fit": self.culture_fit / total,
        }

class RecruitmentState(TypedDict):
    #Job
    raw_jd: str
    job_profile: JobProfile

    #Candidate
    candidate_id: str
    file_path: str
    s3_key: str
    resume_text: str
    candidate_profile: CandidateProfile

    #Matching
    match_score:float
    scoring_weights: dict
    shortlisted:bool
    
    #Interview
    interview_format: str
    interview_status: str
    interview_details: dict
    cancel_reason: str
    transcript_s3_key: str
    manual_scorecard: dict
    interview_scorecard: dict
    match_assessment: dict

    #Final
    recommendation:dict
    decision:str

    #Meta
    job_id: str

#Define graph
graph=StateGraph(RecruitmentState)

def jd_structuring(state: RecruitmentState) -> dict:
    raw_jd = state["raw_jd"]
    if len(raw_jd.strip()) < 100:
        raise ValueError("Job description is too short to be valid. Please provide a complete JD.")
    messages = [
        SystemMessage(content="You are an expert ATS system that extracts structured data from job descriptions. Return only valid JSON."),
        HumanMessage(content=f"""
Extract structured data from this Job Description:

{raw_jd}

Return JSON with this schema:

{{
  "job_title": "",
  "company": "",
  "location": "",
  "employment_type": "",
  "seniority_level": "",
  "domain": "",

  "experience_required": {{
    "min_years": null,
    "max_years": null,
    "raw": ""
  }},

  "skills": {{
    "must_have": [],
    "nice_to_have": []
  }},

  "responsibilities": [],
  "qualifications": [],
  "education": [],

  "keywords": [],

  "screening_questions": [
    {{
      "question": "",
      "evaluates": ""
    }}
  ],

  "signals": {{
    "remote_allowed": null,
    "startup_environment": null,
    "client_facing": null
  }}
}}
""")
    ]
    response = llm.invoke(messages)
    return {"job_profile": JobProfile.model_validate(_parse_llm_json(response.content))}

def ingest_resume(uploaded_file_path: str, raw_jd: str, weights: dict = None, interview_format: str = "video", job_id: str = None) -> dict:
    allowed_extensions = [".pdf", ".docx"]
    extension = os.path.splitext(uploaded_file_path)[1].lower()

    if extension not in allowed_extensions:
        raise ValueError(f"Unsupported file type: {extension}")

    valid_formats = {"video", "phone", "onsite"}
    if interview_format not in valid_formats:
        raise ValueError(f"interview_format must be one of {valid_formats}. Got: '{interview_format}'")

    max_bytes = 10 * 1024 * 1024  # 10 MB
    if os.path.getsize(uploaded_file_path) > max_bytes:
        raise ValueError("Resume exceeds 10MB limit.")

    # Duplicate detection — hash file content, check S3 for existing entry
    with open(uploaded_file_path, "rb") as f:
        file_hash = hashlib.sha256(f.read()).hexdigest()
    dedup_key = f"dedup/{file_hash}"

    from botocore.exceptions import ClientError
    try:
        obj = s3.get_object(Bucket=S3_BUCKET, Key=dedup_key)
        existing_id = obj["Body"].read().decode()
        logger.info(f"Duplicate resume detected — returning existing candidate_id={existing_id}")
        return {
            "candidate_id": existing_id,
            "s3_key": f"resumes/{existing_id}{extension}",
            "status": "DUPLICATE",
        }
    except ClientError as e:
        if e.response["Error"]["Code"] != "NoSuchKey":
            raise

    candidate_id = str(uuid.uuid4())
    s3_key = f"resumes/{candidate_id}{extension}"

    try:
        s3.upload_file(uploaded_file_path, S3_BUCKET, s3_key)
        s3.put_object(Bucket=S3_BUCKET, Key=dedup_key, Body=candidate_id.encode())
    except Exception as e:
        raise RuntimeError(f"Failed to upload resume to S3: {e}")

    try:
        sqs.send_message(
            QueueUrl=SQS_QUEUE_URL,
            MessageBody=json.dumps({
                "candidate_id":    candidate_id,
                "s3_key":          s3_key,
                "raw_jd":          raw_jd,
                "scoring_weights": weights or {},
                "interview_format":interview_format,
                "job_id":          job_id or "",
            })
        )
    except Exception as e:
        raise RuntimeError(f"Failed to queue candidate to SQS: {e}")

    return {
        "candidate_id": candidate_id,
        "s3_key": s3_key,
        "status": "QUEUED"
    }

def extract_resume_text(file_path: str) -> str:
    if file_path.lower().endswith(".pdf"):
        try:
            chunks = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        chunks.append(page_text)
                    for table in page.extract_tables():
                        for row in table:
                            row_text = " | ".join(cell for cell in row if cell and cell.strip())
                            if row_text:
                                chunks.append(row_text)
            text = "\n".join(chunks)
        except Exception as e:
            raise ValueError(f"Failed to extract PDF: {e}")
        if not text.strip():
            raise ValueError("PDF appears to be empty or image-only — no text could be extracted.")
        return text

    elif file_path.lower().endswith(".docx"):
        try:
            doc = Document(file_path)
            chunks = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                    if row_text:
                        chunks.append(row_text)
        except Exception as e:
            raise ValueError(f"Failed to extract DOCX: {e}")
        text = "\n".join(chunks)
        if not text.strip():
            raise ValueError("DOCX appears to be empty — no text could be extracted.")
        return text

    raise ValueError("Unsupported file type")

def clean_resume_text(text: str) -> str:
    """
    Normalize extracted resume text.
    """

    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"Page \d+ of \d+", "", text, flags=re.I)

    return text.strip()

def _regex_email(text: str) -> str | None:
    match = re.search(r"[\w.+-]+@[\w-]+\.[a-z]{2,}", text, re.I)
    return match.group(0) if match else None


def _regex_phone(text: str) -> str | None:
    match = re.search(r"[\+\d][\d\s\-().]{7,15}\d", text)
    return match.group(0) if match else None


def parse_resume_llm(resume_text: str) -> dict:
    schema = """
    {
      "name": null,
      "email": null,
      "phone": null,
      "location": null,
      "skills": [],
      "experience": [
        {"company": "", "role": "", "duration": "", "description": ""}
      ],
      "education": [
        {"institution": "", "degree": "", "year": null}
      ],
      "certifications": []
    }
    """
    response = llm.invoke([
        SystemMessage(content=(
            "You are a resume parser. "
            "Extract only explicit information. Never infer missing values. "
            "Return valid JSON only."
        )),
        HumanMessage(content=f"Parse this resume.\n\nSchema:\n{schema}\n\nResume:\n{resume_text}")
    ])
    return _parse_llm_json(response.content)


def _retry_missing_sections(parsed: dict, resume_text: str) -> dict:
    for _attempt in range(2):
        missing = [f for f in ("skills", "experience") if not parsed.get(f)]
        if not missing:
            break
        response = llm.invoke([
            SystemMessage(content="You are a resume parser. Return only valid JSON."),
            HumanMessage(content=f"""
The previous parse returned empty values for: {missing}.
Re-read the resume carefully and extract only those sections.

Resume:
{resume_text}

Return JSON with only these keys: {missing}
""")
        ])
        retry = _parse_llm_json(response.content)
        for field in missing:
            if retry.get(field):
                parsed[field] = retry[field]
    return parsed


def resume_parsing(state: RecruitmentState) -> dict:
    cid = state["candidate_id"]
    logger.info(f"[{cid}] resume_parsing — extracting text from {state['file_path']}")
    raw_text = extract_resume_text(state["file_path"])
    cleaned_text = clean_resume_text(raw_text)

    parsed = parse_resume_llm(cleaned_text)

    # Regex fallback for email (try cleaned text, then raw text)
    if not parsed.get("email"):
        parsed["email"] = _regex_email(cleaned_text) or _regex_email(raw_text)
        if parsed["email"]:
            logger.info(f"[{cid}] resume_parsing — email recovered via regex")
    if not parsed.get("phone"):
        parsed["phone"] = _regex_phone(cleaned_text)

    # Up to 2 LLM retries for missing skills / experience
    parsed = _retry_missing_sections(parsed, cleaned_text)

    # Log warnings for non-critical missing fields but do NOT block the pipeline.
    # name/skills/experience being absent is unusual but not a showstopper —
    # matching still runs, and interview_scheduling catches a missing email.
    for field in ("name", "skills", "experience"):
        if not parsed.get(field):
            logger.warning(f"[{cid}] resume_parsing — {field} missing, continuing anyway")

    # Only pause for email: without it we cannot send the interview invite.
    if not parsed.get("email"):
        logger.warning(f"[{cid}] resume_parsing — email missing, pausing for manual input")
        raise NodeInterrupt(
            "Resume parsing could not find an email address. "
            "Please add the candidate's email via the Actions panel and resume."
        )

    logger.info(
        f"[{cid}] resume_parsing — name={parsed.get('name')!r}, "
        f"skills={len(parsed.get('skills') or [])}, "
        f"experience={len(parsed.get('experience') or [])}"
    )
    return {
        "candidate_profile": CandidateProfile.model_validate(parsed),
        "resume_text": cleaned_text,
    }

def matching_node(state: RecruitmentState) -> dict:
    cid = state["candidate_id"]
    logger.info(f"[{cid}] matching_node — scoring candidate")
    raw_weights = state.get("scoring_weights") or {}
    normalized = ScoringWeights(**raw_weights).normalized()

    criteria = "\n".join(f"- {k}" for k in normalized)

    response = llm.invoke([
        SystemMessage(content=(
            "You are a resume evaluator. Score each dimension independently on a 0-10 scale. "
            "Evaluate only skills, experience, and education. "
            "Do not consider name, location, gender, nationality, age, or any protected characteristic. "
            "Return only valid JSON."
        )),
        HumanMessage(content=f"""
Score this candidate against the job on these criteria:
{criteria}

For each criterion, return a score (0-10) and one sentence of evidence.

Job: {state["job_profile"].model_dump_json()}
Candidate: {state["candidate_profile"].model_dump_json()}

Return JSON like:
{{
  "skills": 7,
  "skills_evidence": "...",
  "experience": 6,
  "experience_evidence": "..."
}}
""")
    ])

    scores = _parse_llm_json(response.content)
    for k in normalized:
        scores[k] = max(0, min(10, scores.get(k, 0)))
    final_score = sum(normalized[k] * scores[k] for k in normalized)
    final_score = round(final_score, 2)
    logger.info(f"[{cid}] matching_node — match_score={final_score}")

    return {
        "match_score": final_score,
        "match_assessment": scores,
    }

def shortlisting_node(state: RecruitmentState) -> dict:
    cid = state["candidate_id"]
    shortlisted = state["match_score"] >= 6.0
    logger.info(f"[{cid}] shortlisting_node — shortlisted={shortlisted} (score={state['match_score']})")
    return {"shortlisted": shortlisted}

def _get_calendar_service():
    from googleapiclient.discovery import build
    from google.oauth2 import service_account
    # Prefer JSON content as env var (works on Railway without file mounts)
    key_content = os.getenv("GOOGLE_SERVICE_ACCOUNT_JSON_CONTENT")
    key_path    = os.getenv("GOOGLE_SERVICE_ACCOUNT_JSON")
    if key_content:
        creds = service_account.Credentials.from_service_account_info(
            json.loads(key_content),
            scopes=["https://www.googleapis.com/auth/calendar"],
        )
    elif key_path:
        creds = service_account.Credentials.from_service_account_file(
            key_path,
            scopes=["https://www.googleapis.com/auth/calendar"],
        )
    else:
        raise RuntimeError(
            "Google Calendar not configured. "
            "Set GOOGLE_SERVICE_ACCOUNT_JSON_CONTENT (paste the JSON key) "
            "or GOOGLE_SERVICE_ACCOUNT_JSON (path to key file)."
        )
    return build("calendar", "v3", credentials=creds)


def _find_free_slot(service, duration_minutes: int) -> datetime:
    calendar_id  = os.getenv("INTERVIEW_CALENDAR_ID", "primary")
    search_start = datetime.utcnow() + timedelta(days=3)
    search_end   = search_start + timedelta(days=7)

    result = service.freebusy().query(body={
        "timeMin": search_start.isoformat() + "Z",
        "timeMax": search_end.isoformat() + "Z",
        "items": [{"id": calendar_id}],
    }).execute()

    busy_periods = [
        (
            datetime.fromisoformat(b["start"].replace("Z", "")),
            datetime.fromisoformat(b["end"].replace("Z", "")),
        )
        for b in result["calendars"][calendar_id]["busy"]
    ]

    candidate = search_start.replace(hour=9, minute=0, second=0, microsecond=0)
    while candidate < search_end:
        if candidate.weekday() >= 5:
            candidate += timedelta(days=1)
            candidate = candidate.replace(hour=9, minute=0, second=0, microsecond=0)
            continue
        if candidate.hour >= 17:
            candidate += timedelta(days=1)
            candidate = candidate.replace(hour=9, minute=0, second=0, microsecond=0)
            continue

        slot_end = candidate + timedelta(minutes=duration_minutes)
        conflict = False
        for busy_start, busy_end in busy_periods:
            if candidate < busy_end and slot_end > busy_start:
                conflict = True
                candidate = busy_end
                break

        if not conflict:
            return candidate

    raise RuntimeError("No available interview slots found in the next 10 days.")


def create_calendar_event(
    candidate_email: str,
    interviewer_email: str,
    duration_minutes: int,
) -> dict:
    calendar_id  = os.getenv("INTERVIEW_CALENDAR_ID", "primary")
    service      = _get_calendar_service()
    scheduled_at = _find_free_slot(service, duration_minutes)
    end_at       = scheduled_at + timedelta(minutes=duration_minutes)

    event = service.events().insert(
        calendarId=calendar_id,
        body={
            "summary": f"Interview — {candidate_email}",
            "start": {"dateTime": scheduled_at.isoformat() + "Z", "timeZone": "UTC"},
            "end":   {"dateTime": end_at.isoformat() + "Z",       "timeZone": "UTC"},
            "attendees": [
                {"email": candidate_email},
                {"email": interviewer_email},
            ],
            "conferenceData": {
                "createRequest": {
                    "requestId": str(uuid.uuid4()),
                    "conferenceSolutionKey": {"type": "hangoutsMeet"},
                }
            },
        },
        conferenceDataVersion=1,
        sendUpdates="all",
    ).execute()

    meet_link = (
        event.get("conferenceData", {})
            .get("entryPoints", [{}])[0]
            .get("uri", "")
    )
    return {
        "calendar_event_id": event["id"],
        "scheduled_at":      scheduled_at.isoformat(),
        "meeting_link":      meet_link,
    }

def _ses_send(to: str, subject: str, body: str) -> None:
    resend_key = os.getenv("RESEND_API_KEY")
    if resend_key:
        try:
            import resend as _resend
            _resend.api_key = resend_key
            # RESEND_FROM_EMAIL = verified custom domain address (optional)
            # Default: onboarding@resend.dev — works on free plan, no domain setup needed
            from_addr = os.getenv("RESEND_FROM_EMAIL", "onboarding@resend.dev")
            _resend.Emails.send({"from": from_addr, "to": to, "subject": subject, "text": body})
            logger.info(f"Email sent via Resend from={from_addr} -> {to} | {subject}")
            return
        except Exception as e:
            logger.warning(f"Resend failed ({e}) — falling back to SES")

    # Fallback: AWS SES
    try:
        ses.send_email(
            Source=HR_EMAIL,
            Destination={"ToAddresses": [to]},
            Message={
                "Subject": {"Data": subject},
                "Body":    {"Text": {"Data": body}},
            },
        )
        logger.info(f"Email sent via SES -> {to} | {subject}")
    except Exception as e:
        logger.warning(
            f"Email could not be sent (no Resend key + SES sandbox) -> {to}\n"
            f"  Subject : {subject}\n"
            f"  Body    :\n{body}"
        )


def send_candidate_invite(candidate: CandidateProfile, details: dict) -> None:
    formatted_date = datetime.fromisoformat(details["scheduled_at"]).strftime(
        "%A, %d %B %Y at %I:%M %p UTC"
    )
    _ses_send(
        to=candidate.email,
        subject=f"Interview Invitation — {details['job_title']} at {details['company']}",
        body=(
            f"Dear {candidate.name},\n\n"
            f"We are pleased to inform you that following a review of your application, "
            f"you have been selected for an interview for the position of "
            f"{details['job_title']} at {details['company']}.\n\n"
            f"Your interview has been scheduled as follows:\n\n"
            f"  Date & Time  : {formatted_date}\n"
            f"  Format       : {details['format'].capitalize()}\n"
            f"  Meeting Link : {details['meeting_link']}\n\n"
            f"Please join a few minutes early to ensure a smooth start. "
            f"If you are unable to attend or have any questions, "
            f"please reach out to us at {HR_EMAIL}.\n\n"
            f"We look forward to speaking with you.\n\n"
            f"Warm regards,\n"
            f"The Talent Acquisition Team\n"
            f"{details['company']}"
        ),
    )

def send_interviewer_brief(
    interviewer_email: str,
    candidate: CandidateProfile,
    match_score: float,
    match_assessment: dict,
    resume_url: str,
) -> None:
    dimensions = ["skills", "experience", "education", "culture_fit"]
    assessment_lines = "\n".join(
        f"  {dim.replace('_', ' ').title():<14}: {match_assessment.get(dim, 'N/A')}/10"
        f" — {match_assessment.get(f'{dim}_evidence', '')}"
        for dim in dimensions
    )
    _ses_send(
        to=interviewer_email,
        subject=f"Interview Brief — {candidate.name}",
        body=(
            f"Hi,\n\n"
            f"Please find below the brief for your upcoming interview.\n\n"
            f"{'=' * 52}\n"
            f"Candidate    : {candidate.name}\n"
            f"Email        : {candidate.email}\n"
            f"Match Score  : {match_score} / 10\n\n"
            f"AI Assessment Breakdown:\n"
            f"{assessment_lines}\n\n"
            f"Resume       : {resume_url}\n"
            f"{'=' * 52}\n\n"
            f"Kindly review the candidate's profile before the session.\n\n"
            f"Regards,\n"
            f"The Talent Acquisition Team"
        ),
    )


def interview_scheduling(state: RecruitmentState) -> dict:
    cid = state["candidate_id"]
    candidate = state["candidate_profile"]
    logger.info(f"[{cid}] interview_scheduling — scheduling for {candidate.name}")

    if not candidate.email:
        raise NodeInterrupt(
            f"Cannot schedule interview for {candidate.name or 'Unknown'}: "
            f"no email address on resume. Add it via the Actions panel and resume."
        )

    resume_url = s3.generate_presigned_url(
        "get_object",
        Params={"Bucket": S3_BUCKET, "Key": state["s3_key"]},
        ExpiresIn=7 * 24 * 3600,  # 7 days
    )

    interview_id      = str(uuid.uuid4())
    interviewer_email = os.getenv("INTERVIEWER_EMAIL", "hiring@company.com")

    try:
        calendar = create_calendar_event(
            candidate_email=candidate.email,
            interviewer_email=interviewer_email,
            duration_minutes=15,
        )
        logger.info(f"[{cid}] Google Calendar event created — {calendar['scheduled_at']}")
    except Exception as cal_err:
        # Fallback: schedule 2 business days from now, use DEFAULT_MEETING_LINK
        logger.warning(f"[{cid}] Google Calendar unavailable ({cal_err}) — using fallback slot")
        fallback_dt = (datetime.utcnow() + timedelta(days=2)).replace(
            hour=10, minute=0, second=0, microsecond=0
        )
        # Skip Saturday (5) and Sunday (6)
        while fallback_dt.weekday() >= 5:
            fallback_dt += timedelta(days=1)
        calendar = {
            "calendar_event_id": None,
            "scheduled_at":      fallback_dt.isoformat(),
            "meeting_link":      os.getenv("DEFAULT_MEETING_LINK", "https://meet.google.com/placeholder"),
        }

    details = {
        "interview_id":      interview_id,
        "scheduled_at":      calendar["scheduled_at"],
        "format":            state["interview_format"],
        "duration_minutes":  15,
        "meeting_link":      calendar["meeting_link"],
        "interviewer_email": interviewer_email,
        "resume_url":        resume_url,
        "calendar_event_id": calendar.get("calendar_event_id"),
        "job_title":         state["job_profile"].job_title,
        "company":           state["job_profile"].company,
    }

    send_candidate_invite(candidate, details)
    send_interviewer_brief(
        interviewer_email=interviewer_email,
        candidate=candidate,
        match_score=state["match_score"],
        match_assessment=state["match_assessment"],
        resume_url=resume_url,
    )

    logger.info(f"[{cid}] interview_scheduling — scheduled at {calendar['scheduled_at']}, format={state['interview_format']}")
    return {
        "interview_status": "scheduled",
        "interview_details": details,
    }

def wait_for_interview(state: RecruitmentState) -> dict:
    cid = state["candidate_id"]
    if state.get("cancel_reason"):
        logger.info(f"[{cid}] wait_for_interview — cancelled: {state['cancel_reason']}")
        return {"interview_status": "cancelled", "decision": "rejected"}

    # If recruiter explicitly marked complete, trust that and proceed immediately
    # regardless of scheduled time (allows demo / early completion)
    if state.get("interview_status") == "completed":
        logger.info(f"[{cid}] wait_for_interview — marked complete by recruiter, proceeding")
        return {"interview_status": "completed"}

    scheduled_at = datetime.fromisoformat(state["interview_details"]["scheduled_at"])
    now = datetime.utcnow()

    if now < scheduled_at:
        logger.info(f"[{cid}] wait_for_interview — pausing until {scheduled_at.isoformat()}")
        raise NodeInterrupt(
            f"Interview scheduled for {scheduled_at.strftime('%A, %d %B %Y at %I:%M %p UTC')}. "
            f"Resume this graph once the interview is completed."
        )

    logger.info(f"[{cid}] wait_for_interview — time passed, awaiting confirmation")
    raise NodeInterrupt(
        "Interview time has passed. Awaiting interviewer confirmation to proceed."
    )

def _build_evaluation_dimensions(job: JobProfile) -> list:
    core = [
        {"name": "Communication",   "weight": 0.15},
        {"name": "Problem Solving", "weight": 0.20},
        {"name": "Culture Fit",     "weight": 0.15},
    ]
    skills = job.skills.must_have[:5]
    skill_weight = round(0.50 / max(len(skills), 1), 2)
    technical = [{"name": s, "weight": skill_weight} for s in skills]
    return core + technical


def generate_scorecard_from_transcript(transcript: str, job: JobProfile) -> dict:
    dimensions = _build_evaluation_dimensions(job)
    dimension_names = [d["name"] for d in dimensions]

    response = llm.invoke([
        SystemMessage(content=(
            "You are a senior technical interviewer scoring a candidate. "
            "Be objective and evidence-based. "
            "Evaluate only technical performance and communication quality. "
            "Do not factor in accent, name, nationality, gender, or any protected characteristic. "
            "Return only valid JSON."
        )),
        HumanMessage(content=f"""
Evaluate this interview for the role of {job.job_title} ({job.seniority_level}) at {job.company}.

Dimensions to score: {dimension_names}

Transcript:
{transcript}

Return JSON:
{{
  "dimensions": {{
    "<dimension_name>": {{
      "score": 4,
      "evidence": "specific quote or observation from the transcript"
    }}
  }},
  "strengths": ["..."],
  "concerns":  ["..."],
  "recommendation": "strong_yes|yes|no|strong_no",
  "summary": "2-3 sentence overall assessment"
}}

Rules:
- Score 1-5 (1=poor, 3=meets expectations, 5=exceptional)
- Evidence must be specific — quote or paraphrase from the transcript
- Recommendation: strong_yes (>4.0), yes (>3.0), no (>2.0), strong_no (≤2.0)
- Return a key for every dimension in the list above
""")
    ])

    raw = _parse_llm_json(response.content)
    scored = raw.get("dimensions", {})
    total_weight = sum(d["weight"] for d in dimensions)
    overall = sum(
        max(1, min(5, scored.get(d["name"], {}).get("score", 0))) * d["weight"]
        for d in dimensions
    ) / total_weight

    valid_recommendations = {"strong_yes", "yes", "no", "strong_no"}
    recommendation = raw.get("recommendation", "")
    if recommendation not in valid_recommendations:
        recommendation = "no"

    return {
        "dimensions": [
            {
                "name":     d["name"],
                "weight":   d["weight"],
                "score":    max(1, min(5, scored.get(d["name"], {}).get("score", 1))),
                "evidence": scored.get(d["name"], {}).get("evidence", ""),
            }
            for d in dimensions
        ],
        "overall_score":    round(overall, 2),
        "strengths":        raw.get("strengths", []),
        "concerns":         raw.get("concerns", []),
        "recommendation":   recommendation,
        "summary":          raw.get("summary", ""),
    }


def scorecard_ingestion(state: RecruitmentState) -> dict:
    cid = state["candidate_id"]
    if state.get("interview_status") == "cancelled":
        logger.info(f"[{cid}] scorecard_ingestion — skipped (cancelled)")
        return {"interview_scorecard": {}, "decision": "rejected"}

    interview_format = state["interview_format"]
    logger.info(f"[{cid}] scorecard_ingestion — format={interview_format}")

    if interview_format in ("video", "phone"):
        transcript_s3_key = state.get("transcript_s3_key")
        if not transcript_s3_key:
            logger.info(f"[{cid}] scorecard_ingestion — pausing, transcript not yet uploaded")
            raise NodeInterrupt(
                "Waiting for interview transcript. "
                "Resume with transcript_s3_key once the recording is processed."
            )
        fd, tmp_path = tempfile.mkstemp(suffix=".txt")
        os.close(fd)
        try:
            s3.download_file(S3_BUCKET, transcript_s3_key, tmp_path)
        except Exception as e:
            os.remove(tmp_path)
            raise RuntimeError(f"Failed to download transcript from S3: {e}")
        try:
            with open(tmp_path, "r", encoding="utf-8", errors="replace") as f:
                transcript = f.read()
        finally:
            os.remove(tmp_path)

        scorecard = generate_scorecard_from_transcript(transcript, state["job_profile"])
        logger.info(f"[{cid}] scorecard_ingestion — overall_score={scorecard.get('overall_score')}, recommendation={scorecard.get('recommendation')}")

    else:  # onsite — human fills scorecard manually
        manual_scorecard = state.get("manual_scorecard")
        if not manual_scorecard:
            logger.info(f"[{cid}] scorecard_ingestion — pausing, awaiting manual scorecard")
            raise NodeInterrupt(
                "Onsite interview completed. "
                "Resume with manual_scorecard to proceed."
            )
        scorecard = manual_scorecard
        logger.info(f"[{cid}] scorecard_ingestion — manual scorecard received")

    return {"interview_scorecard": scorecard}

def recommendation_node(state: RecruitmentState) -> dict:
    cid = state["candidate_id"]
    if state.get("decision") == "rejected":
        logger.info(f"[{cid}] recommendation_node — skipped (already rejected)")
        return {"recommendation": {
            "hire_recommendation": "no",
            "summary": "Interview was cancelled or candidate was rejected.",
        }}

    logger.info(f"[{cid}] recommendation_node — synthesizing pre+post signals")
    pre_score = state["match_score"]                          # 0-10
    scorecard = state.get("interview_scorecard", {})
    raw_interview_score = scorecard.get("overall_score", 0)
    post_score = round((raw_interview_score / 5) * 10, 2)    # normalize 1-5 → 0-10

    final_score = round(0.4 * pre_score + 0.6 * post_score, 2)

    gap = abs(pre_score - post_score)
    if gap < 1.5:
        alignment, confidence = "strong", "high"
    elif gap < 3.0:
        alignment, confidence = "moderate", "medium"
    else:
        alignment, confidence = "discrepant", "low"

    response = llm.invoke([
        SystemMessage(content=(
            "You are a hiring manager making a final recommendation. "
            "Be concise and objective. Return only valid JSON."
        )),
        HumanMessage(content=f"""
Synthesize these two assessments for {state["candidate_profile"].name}
applying for {state["job_profile"].job_title} at {state["job_profile"].company}.

PRE-INTERVIEW AI ASSESSMENT (40% weight):
Score       : {pre_score}/10
Assessment  : {json.dumps(state["match_assessment"], indent=2)}

POST-INTERVIEW EVALUATION (60% weight):
Score       : {post_score}/10 (normalized from {raw_interview_score}/5)
Strengths   : {scorecard.get("strengths", [])}
Concerns    : {scorecard.get("concerns", [])}
Interviewer : {scorecard.get("recommendation", "")}
Summary     : {scorecard.get("summary", "")}

Alignment   : {alignment} (gap of {round(gap, 1)} points between pre and post)
Final Score : {final_score}/10

Return JSON:
{{
  "hire_recommendation": "strong_yes|yes|no|strong_no",
  "summary": "3-4 sentences synthesizing both assessments and explaining the recommendation",
  "strengths": [],
  "concerns": [],
  "flag_for_human": true
}}

Rules:
- flag_for_human = true if alignment is discrepant OR final_score is between 4.5-6.5
- Be specific — reference actual evidence from both assessments
- If discrepant, explain which signal to trust more and why
""")
    ])

    raw = _parse_llm_json(response.content)

    valid_recs = {"strong_yes", "yes", "no", "strong_no"}
    hire_rec = raw.get("hire_recommendation", "no")
    if hire_rec not in valid_recs:
        hire_rec = "no"

    logger.info(f"[{cid}] recommendation_node — final_score={final_score}, recommendation={hire_rec}, alignment={alignment}, flag_for_human={raw.get('flag_for_human', alignment == 'discrepant')}")
    return {
        "recommendation": {
            "final_score":          final_score,
            "pre_interview_score":  pre_score,
            "post_interview_score": post_score,
            "hire_recommendation":  hire_rec,
            "confidence":           confidence,
            "alignment":            alignment,
            "strengths":            raw.get("strengths", []),
            "concerns":             raw.get("concerns", []),
            "summary":              raw.get("summary", ""),
            "flag_for_human":       raw.get("flag_for_human", alignment == "discrepant"),
        }
    }

def _send_offer_notification(candidate: CandidateProfile, _recommendation: dict) -> None:
    _ses_send(
        to=candidate.email,
        subject="Congratulations — We'd Like to Move Forward",
        body=(
            f"Dear {candidate.name},\n\n"
            f"We are thrilled to inform you that after a thorough review of your profile "
            f"and interview performance, our team has decided to move forward with you.\n\n"
            f"Our HR team will be in touch shortly with the formal offer details and next steps.\n\n"
            f"We look forward to welcoming you to the team.\n\n"
            f"Warm regards,\n"
            f"The Talent Acquisition Team"
        ),
    )

def _send_rejection_notification(candidate: CandidateProfile) -> None:
    _ses_send(
        to=candidate.email,
        subject="Update on Your Application",
        body=(
            f"Dear {candidate.name},\n\n"
            f"Thank you for the time and effort you invested in our interview process. "
            f"We truly appreciate your interest in joining our team.\n\n"
            f"After careful consideration, we have decided to move forward with other "
            f"candidates whose experience more closely matches our current requirements.\n\n"
            f"We encourage you to apply for future openings that align with your background.\n\n"
            f"We wish you all the best in your search.\n\n"
            f"Warm regards,\n"
            f"The Talent Acquisition Team"
        ),
    )

def human_approval(state: RecruitmentState) -> dict:
    cid = state["candidate_id"]
    decision = state.get("decision")

    if decision in ("approved", "rejected"):
        candidate = state["candidate_profile"]
        recommendation = state.get("recommendation", {})
        logger.info(f"[{cid}] human_approval — decision={decision} for {candidate.name}")
        if decision == "approved":
            _send_offer_notification(candidate, recommendation)
        else:
            _send_rejection_notification(candidate)
        return {"decision": decision}

    recommendation = state.get("recommendation", {})
    flag = recommendation.get("flag_for_human", False)
    candidate = state["candidate_profile"]
    job = state["job_profile"]

    if flag:
        interrupt_msg = (
            f"REVIEW REQUIRED — {candidate.name} | {job.job_title} at {job.company}\n"
            f"{'─' * 55}\n"
            f"Pre-Interview Score  : {recommendation.get('pre_interview_score')}/10\n"
            f"Post-Interview Score : {recommendation.get('post_interview_score')}/10\n"
            f"Final Blended Score  : {recommendation.get('final_score')}/10\n"
            f"Alignment            : {recommendation.get('alignment')} "
            f"(confidence: {recommendation.get('confidence')})\n\n"
            f"Strengths : {recommendation.get('strengths', [])}\n"
            f"Concerns  : {recommendation.get('concerns', [])}\n\n"
            f"AI Recommendation : {recommendation.get('hire_recommendation')}\n"
            f"Summary           : {recommendation.get('summary')}\n"
            f"{'─' * 55}\n"
            f"Resume with decision='approved' or decision='rejected'."
        )
    else:
        interrupt_msg = (
            f"Final Approval — {candidate.name} | {job.job_title}\n"
            f"Score: {recommendation.get('final_score')}/10  "
            f"Recommendation: {recommendation.get('hire_recommendation')}\n"
            f"Summary: {recommendation.get('summary')}\n\n"
            f"Resume with decision='approved' or decision='rejected'."
        )

    logger.info(f"[{cid}] human_approval — pausing for review (flag_for_human={flag})")
    raise NodeInterrupt(interrupt_msg)

def shortlist_rejection(state: RecruitmentState) -> dict:
    cid = state["candidate_id"]
    candidate = state["candidate_profile"]
    job = state["job_profile"]
    _ses_send(
        to=candidate.email,
        subject=f"Your Application at {job.company}",
        body=(
            f"Dear {candidate.name},\n\n"
            f"Thank you for your interest in the {job.job_title} role at {job.company}.\n\n"
            f"After carefully reviewing your application, we have decided to move forward "
            f"with candidates whose experience more closely matches our current requirements.\n\n"
            f"We appreciate the time you took to apply and encourage you to apply for "
            f"future openings that align with your background.\n\n"
            f"We wish you all the best.\n\n"
            f"Warm regards,\n"
            f"The Talent Acquisition Team\n"
            f"{job.company}"
        ),
    )
    logger.info(f"[{cid}] shortlist_rejection — rejection sent to {candidate.email} for {job.job_title}")
    return {"decision": "rejected"}


def route_after_shortlist(state: RecruitmentState) -> str:
    return "Interview_scheduling" if state["shortlisted"] else "Shortlist_rejection"

#Add nodes
graph.add_node("JD_structuring", jd_structuring)
graph.add_node("Resume_parsing", resume_parsing)
graph.add_node("Matching_node", matching_node)
graph.add_node("Shortlisting_node", shortlisting_node)
graph.add_node("Interview_scheduling", interview_scheduling)
graph.add_node("WAIT_FOR_INTERVIEW", wait_for_interview)
graph.add_node("Scorecard_Ingestion", scorecard_ingestion)
graph.add_node("Recommendation_node", recommendation_node)
graph.add_node("Human_approval", human_approval)
graph.add_node("Shortlist_rejection", shortlist_rejection)

#Wire edges
graph.add_edge(START, "JD_structuring")
graph.add_edge("JD_structuring","Resume_parsing")
graph.add_edge("Resume_parsing", "Matching_node")
graph.add_edge("Matching_node", "Shortlisting_node")
graph.add_conditional_edges("Shortlisting_node", route_after_shortlist)
graph.add_edge("Interview_scheduling", "WAIT_FOR_INTERVIEW")
graph.add_edge("WAIT_FOR_INTERVIEW", "Scorecard_Ingestion")
graph.add_edge("Scorecard_Ingestion", "Recommendation_node")
graph.add_edge("Recommendation_node", "Human_approval")
graph.add_edge("Human_approval", END)
graph.add_edge("Shortlist_rejection", END)

_postgres_url = os.getenv("DATABASE_URL") or os.getenv("POSTGRES_URL")
if _postgres_url:
    from langgraph.checkpoint.postgres import PostgresSaver
    checkpointer = PostgresSaver.from_conn_string(_postgres_url)
    checkpointer.setup()
    _conn = None
else:
    import sqlite3
    from langgraph.checkpoint.sqlite import SqliteSaver
    _db_path = os.getenv("DATABASE_PATH", "hiring_agent.db")
    _conn = sqlite3.connect(_db_path, check_same_thread=False)
    checkpointer = SqliteSaver(_conn)

graph = graph.compile(checkpointer=checkpointer)


